#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
render_urteil.py - Generiert deutsche Zivilurteile im offiziellen Gerichtslayout.

Unterstuetzte Dokumenttypen:
  - urteil        Vollurteil nach Paragraf 313 ZPO (Rubrum, Tenor, Tatbestand,
                  Entscheidungsgruende, Rechtsmittelbelehrung, Unterschriften)
  - versaeumnis   Versaeumnisurteil (kein Tatbestand, keine Gruende)
  - beschluss     Beschluss (Rubrum, Tenor, Gruende kurz, Rechtsmittel)
  - relation      Relationsgutachten im Schul-Layout (kein Volksformel, Stempel
                  Relation oben)

Eingabe: Ordner mit
  - rubrum.yaml                 Gericht, Aktenzeichen, Parteien, Termin, Spruchkoerper
  - tenor.md                    Nummerierte Tenor-Punkte
  - tatbestand.md               Tatbestand (Markdown mit **fett**, *kursiv*)
  - entscheidungsgruende.md     Entscheidungsgruende (Markdown)
  - rechtsmittelbelehrung.md    Belehrung (Markdown)
  - relation.md                 (nur fuer typ=relation) komplette Vollrelation

Ausgabe: DOCX (Arial 11pt, 2.5/2/2/2 cm). Optional PDF via soffice.

Layout-Quellen:
  - OLG Sachsen-Anhalt Musterurteil Zivil (Ausbildungsmappe Referendariat 2012)
  - Repetitorium Hofmann Skript Einfuehrung Referendariat ZPO
  - iurratio Relationsgutachten

Autor: Klotzkette - Plugin urteilsbauer-relationsmacher
"""

import argparse
import re
import shutil
import subprocess
import sys
from pathlib import Path

try:
    import yaml
except ImportError:
    yaml = None

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("FEHLER: python-docx fehlt. Bitte 'pip install python-docx pyyaml' ausfuehren.",
          file=sys.stderr)
    sys.exit(2)


# ---------- Schriftkonstanten Gerichtsstandard ----------

SCHRIFT = "Arial"
SCHRIFTGROESSE_NORMAL = Pt(11)
SCHRIFTGROESSE_KLEIN = Pt(9)
SCHRIFTGROESSE_GERICHT = Pt(14)
SCHRIFTGROESSE_TITEL = Pt(16)

# Inline-Markdown Pattern: **fett**, *kursiv*, _kursiv_
_INLINE_PATTERN = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|_[^_]+_)")


# ---------- Lade-Funktionen ----------

def lade_rubrum(eingabe: Path) -> dict:
    """Laedt rubrum.yaml mit Sinnvoll-Defaults fuer Pflichtfelder."""
    defaults = {
        "gericht": "Amtsgericht Musterstadt",
        "aktenzeichen": "11 C 2406/26",
        "verkuendet_am": "01.01.2026",
        "urkundsbeamter": "Schmidt",
        "klaeger": {
            "name": "Mustermann GmbH",
            "anschrift": "Musterstrasse 1 in 12345 Musterstadt",
            "rolle": "Klaegerin",
            "prozessbevollmaechtigte": "Rechtsanwaeltin Dr. Mueller in Hamburg",
        },
        "beklagter": {
            "name": "Beispiel AG",
            "anschrift": "Beispielweg 2 in 67890 Beispielort",
            "rolle": "Beklagte",
            "prozessbevollmaechtigte": "Rechtsanwalt Maier in Frankfurt",
        },
        "spruchkoerper": "die Richterin am Amtsgericht Dr. Beispiel",
        "termin_muendliche_verhandlung": "15. April 2026",
        "schluss_muendlich": "15. April 2026",
        "im_namen_des_volkes": True,
    }
    pfad = eingabe / "rubrum.yaml"
    if not pfad.exists():
        return defaults
    if yaml is None:
        return defaults
    geladen = yaml.safe_load(pfad.read_text(encoding="utf-8")) or {}
    for k, v in defaults.items():
        geladen.setdefault(k, v)
    return geladen


def lies_text(eingabe: Path, dateiname: str) -> str:
    pfad = eingabe / dateiname
    if not pfad.exists():
        return ""
    return pfad.read_text(encoding="utf-8")


# ---------- Layout-Helfer ----------

def setze_standardschrift(dokument: Document, schrift: str) -> None:
    style = dokument.styles["Normal"]
    style.font.name = schrift
    style.font.size = SCHRIFTGROESSE_NORMAL
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), schrift)
    rfonts.set(qn("w:hAnsi"), schrift)
    rfonts.set(qn("w:cs"), schrift)


def setze_seitenraender(dokument: Document) -> None:
    for section in dokument.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)


def _schreibe_inline_text(absatz, text: str, *, fett: bool = False,
                          kursiv: bool = False, groesse=None) -> None:
    if not text:
        return
    teile = _INLINE_PATTERN.split(text)
    for teil in teile:
        if not teil:
            continue
        ist_fett = fett
        ist_kursiv = kursiv
        inhalt = teil
        if teil.startswith("**") and teil.endswith("**") and len(teil) > 4:
            ist_fett = True
            inhalt = teil[2:-2]
        elif teil.startswith("*") and teil.endswith("*") and len(teil) > 2:
            ist_kursiv = True
            inhalt = teil[1:-1]
        elif teil.startswith("_") and teil.endswith("_") and len(teil) > 2:
            ist_kursiv = True
            inhalt = teil[1:-1]
        run = absatz.add_run(inhalt)
        run.bold = ist_fett
        run.italic = ist_kursiv
        if groesse:
            run.font.size = groesse


def absatz_zentriert(dokument, text, *, fett=False, groesse=None, kursiv=False):
    p = dokument.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _schreibe_inline_text(p, text, fett=fett, kursiv=kursiv, groesse=groesse)


def absatz_rechts(dokument, text, *, kursiv=False, groesse=None):
    p = dokument.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _schreibe_inline_text(p, text, kursiv=kursiv, groesse=groesse)


def absatz_links(dokument, text, *, fett=False, kursiv=False, groesse=None):
    p = dokument.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _schreibe_inline_text(p, text, fett=fett, kursiv=kursiv, groesse=groesse)


def absatz_blocksatz(dokument, text):
    p = dokument.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _schreibe_inline_text(p, text)


def leerzeile(dokument):
    dokument.add_paragraph()


# ---------- Bloecke ----------

def schreibe_kopf(dokument, rubrum):
    az = rubrum.get("aktenzeichen", "")
    if az:
        absatz_rechts(dokument, "Geschaefts-Nr.: " + str(az),
                      kursiv=True, groesse=SCHRIFTGROESSE_KLEIN)
    verkuendet = rubrum.get("verkuendet_am", "")
    if verkuendet:
        absatz_rechts(dokument, "Verkuendet am " + str(verkuendet),
                      kursiv=True, groesse=SCHRIFTGROESSE_KLEIN)
    urkund = rubrum.get("urkundsbeamter", "")
    if urkund:
        absatz_rechts(dokument, str(urkund) + ", Urkundsbeamtin der Geschaeftsstelle",
                      kursiv=True, groesse=SCHRIFTGROESSE_KLEIN)
    leerzeile(dokument)


def schreibe_gerichtskopf(dokument, rubrum, *, ohne_volksformel=False):
    gericht = rubrum.get("gericht", "Amtsgericht")
    absatz_zentriert(dokument, str(gericht), fett=True, groesse=SCHRIFTGROESSE_GERICHT)
    leerzeile(dokument)
    if not ohne_volksformel and rubrum.get("im_namen_des_volkes", True):
        absatz_zentriert(dokument, "Im Namen des Volkes", fett=True,
                         groesse=SCHRIFTGROESSE_NORMAL)
        leerzeile(dokument)


def schreibe_urteilstitel(dokument, titel="Urteil"):
    absatz_zentriert(dokument, titel, fett=True, groesse=SCHRIFTGROESSE_TITEL)
    leerzeile(dokument)


def _partei_block(dokument, partei, nummer=None, total=1):
    name = partei.get("name", "")
    anschrift = partei.get("anschrift", "")
    rolle = partei.get("rolle", "")
    prozessb = partei.get("prozessbevollmaechtigte", "")
    praefix = ""
    if nummer is not None and total > 1:
        praefix = str(nummer) + ". "
    absatz_links(dokument, praefix + str(name), fett=True)
    if anschrift:
        absatz_links(dokument, str(anschrift))
    if rolle:
        absatz_links(dokument, "- " + str(rolle) + " -")
    if prozessb:
        absatz_links(dokument, "Prozessbevollmaechtigte: " + str(prozessb))


def schreibe_rubrum(dokument, rubrum):
    absatz_links(dokument, "In dem Rechtsstreit")
    leerzeile(dokument)
    klaeger = rubrum.get("klaeger")
    if isinstance(klaeger, list):
        for i, partei in enumerate(klaeger, 1):
            _partei_block(dokument, partei, nummer=i, total=len(klaeger))
            leerzeile(dokument)
    elif isinstance(klaeger, dict):
        _partei_block(dokument, klaeger)
        leerzeile(dokument)
    absatz_zentriert(dokument, "gegen")
    leerzeile(dokument)
    beklagter = rubrum.get("beklagter")
    if isinstance(beklagter, list):
        for i, partei in enumerate(beklagter, 1):
            _partei_block(dokument, partei, nummer=i, total=len(beklagter))
            leerzeile(dokument)
    elif isinstance(beklagter, dict):
        _partei_block(dokument, beklagter)
        leerzeile(dokument)
    gericht = rubrum.get("gericht", "Amtsgericht")
    spruch = rubrum.get("spruchkoerper", "die zustaendige Richterin")
    termin = rubrum.get("termin_muendliche_verhandlung",
                        rubrum.get("schluss_muendlich", ""))
    einleitung = ("hat das " + str(gericht) + " durch " + str(spruch))
    if termin:
        einleitung += (" auf die muendliche Verhandlung vom " + str(termin))
    einleitung += " fuer Recht erkannt:"
    absatz_blocksatz(dokument, einleitung)
    leerzeile(dokument)


def schreibe_tenor(dokument, text):
    if not text.strip():
        return
    zeilen = [z for z in text.splitlines() if z.strip()]
    punkt_pattern = re.compile(r"^\s*\(?(\d+)\)?[\.\)]\s*(.+)$")
    aktuelle_nummer = 0
    aktueller_text = []

    def flush(nr, text_teile):
        if not text_teile:
            return
        zusammen = " ".join(t.strip() for t in text_teile if t.strip())
        if not zusammen:
            return
        p = dokument.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.first_line_indent = Cm(-0.8)
        _schreibe_inline_text(p, str(nr) + ". " + zusammen)

    for zeile in zeilen:
        treffer = punkt_pattern.match(zeile)
        if treffer:
            flush(aktuelle_nummer, aktueller_text)
            aktuelle_nummer = int(treffer.group(1))
            aktueller_text = [treffer.group(2)]
        else:
            aktueller_text.append(zeile)
    flush(aktuelle_nummer, aktueller_text)
    leerzeile(dokument)


def schreibe_abschnitt(dokument, ueberschrift, text):
    if not text.strip():
        return
    absatz_links(dokument, ueberschrift, fett=True, groesse=SCHRIFTGROESSE_NORMAL)
    leerzeile(dokument)
    aktueller_block = []

    def flush():
        if not aktueller_block:
            return
        zusammen = "\n".join(aktueller_block).strip()
        if not zusammen:
            aktueller_block.clear()
            return
        absatz_blocksatz(dokument, " ".join(
            z.strip() for z in zusammen.splitlines() if z.strip()))
        aktueller_block.clear()

    for zeile in text.splitlines():
        roh = zeile.rstrip()
        if not roh.strip():
            flush()
            continue
        if roh.lstrip().startswith("### "):
            flush()
            absatz_links(dokument, roh.lstrip()[4:], fett=True)
            continue
        if roh.lstrip().startswith("## "):
            flush()
            absatz_links(dokument, roh.lstrip()[3:], fett=True)
            continue
        if roh.lstrip().startswith(("- ", "* ")):
            flush()
            p = dokument.add_paragraph(style="List Bullet")
            _schreibe_inline_text(p, roh.lstrip()[2:])
            continue
        aktueller_block.append(roh)
    flush()
    leerzeile(dokument)


def schreibe_rechtsmittel(dokument, text):
    if not text.strip():
        return
    absatz_links(dokument, "Rechtsmittelbelehrung", fett=True)
    leerzeile(dokument)
    for absatz in text.split("\n\n"):
        absatz = absatz.strip()
        if absatz:
            absatz_blocksatz(dokument, " ".join(
                z.strip() for z in absatz.splitlines() if z.strip()))
    leerzeile(dokument)


def schreibe_unterschrift(dokument, rubrum):
    leerzeile(dokument)
    leerzeile(dokument)
    spruch = rubrum.get("spruchkoerper", "")
    name = ""
    funktion = "Richterin am Amtsgericht"
    if isinstance(spruch, str):
        teile = spruch.split()
        if "Dr." in teile:
            idx = teile.index("Dr.")
            name = " ".join(teile[idx:])
            funktion = " ".join(teile[1:idx]) if len(teile) > 1 else funktion
        else:
            name = teile[-1] if teile else ""
            funktion = " ".join(teile[1:-1]) if len(teile) > 2 else funktion
    absatz_links(dokument, "___________________________")
    absatz_links(dokument, name)
    absatz_links(dokument, funktion, kursiv=True, groesse=SCHRIFTGROESSE_KLEIN)


# ---------- Relations-Modus ----------

def baue_relation(eingabe: Path, ausgabe: Path) -> Path:
    """Rendert ein Relationsgutachten im Schul-Layout (kein Im Namen des Volkes,
    Titel RELATION, Sachbericht und Stationen aus relation.md)."""
    rubrum = lade_rubrum(eingabe)
    relation = lies_text(eingabe, "relation.md")
    dokument = Document()
    setze_standardschrift(dokument, SCHRIFT)
    setze_seitenraender(dokument)
    schreibe_kopf(dokument, rubrum)
    schreibe_gerichtskopf(dokument, rubrum, ohne_volksformel=True)
    schreibe_urteilstitel(dokument, "Relationsgutachten")
    schreibe_rubrum(dokument, rubrum)
    schreibe_abschnitt(dokument, "Relation", relation)
    schreibe_unterschrift(dokument, rubrum)
    ausgabe.parent.mkdir(parents=True, exist_ok=True)
    dokument.save(ausgabe)
    return ausgabe


# ---------- Hauptfunktion ----------

def baue_urteil(eingabe: Path, ausgabe: Path, dokumenttyp: str = "urteil") -> Path:
    if dokumenttyp == "relation":
        return baue_relation(eingabe, ausgabe)
    rubrum = lade_rubrum(eingabe)
    tenor = lies_text(eingabe, "tenor.md")
    tatbestand = lies_text(eingabe, "tatbestand.md")
    gruende = lies_text(eingabe, "entscheidungsgruende.md")
    rechtsmittel = lies_text(eingabe, "rechtsmittelbelehrung.md")

    dokument = Document()
    setze_standardschrift(dokument, SCHRIFT)
    setze_seitenraender(dokument)
    schreibe_kopf(dokument, rubrum)
    schreibe_gerichtskopf(dokument, rubrum)
    titel_map = {
        "urteil": "Urteil",
        "versaeumnis": "Versaeumnisurteil",
        "beschluss": "Beschluss",
    }
    schreibe_urteilstitel(dokument, titel_map.get(dokumenttyp, "Urteil"))
    schreibe_rubrum(dokument, rubrum)
    schreibe_tenor(dokument, tenor)
    if dokumenttyp != "versaeumnis":
        schreibe_abschnitt(dokument, "Tatbestand", tatbestand)
        schreibe_abschnitt(dokument, "Entscheidungsgruende", gruende)
    schreibe_rechtsmittel(dokument, rechtsmittel)
    schreibe_unterschrift(dokument, rubrum)
    ausgabe.parent.mkdir(parents=True, exist_ok=True)
    dokument.save(ausgabe)
    return ausgabe


def konvertiere_zu_pdf(docx_pfad: Path):
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        print("WARNUNG: soffice/libreoffice nicht gefunden, ueberspringe PDF.",
              file=sys.stderr)
        return None
    ausgabe_dir = docx_pfad.parent
    cmd = [soffice, "--headless", "--convert-to", "pdf", "--outdir",
           str(ausgabe_dir), str(docx_pfad)]
    try:
        subprocess.run(cmd, check=True, capture_output=True, timeout=120)
    except subprocess.CalledProcessError as e:
        print("FEHLER bei PDF-Konvertierung: " + str(e), file=sys.stderr)
        return None
    pdf_pfad = ausgabe_dir / (docx_pfad.stem + ".pdf")
    return pdf_pfad if pdf_pfad.exists() else None


def main():
    parser = argparse.ArgumentParser(
        description="Generiert deutsche Zivilurteile im Gerichtslayout.")
    parser.add_argument("eingabe", type=Path,
                        help="Ordner mit rubrum.yaml und den Markdown-Bausteinen")
    parser.add_argument("ausgabe", type=Path, help="Ziel-DOCX-Pfad")
    parser.add_argument("--typ",
                        choices=["urteil", "versaeumnis", "beschluss", "relation"],
                        default="urteil")
    parser.add_argument("--pdf", action="store_true",
                        help="Zusaetzlich PDF erzeugen.")
    args = parser.parse_args()

    if not args.eingabe.exists():
        print("FEHLER: Eingabeordner nicht gefunden: " + str(args.eingabe),
              file=sys.stderr)
        return 1
    docx = baue_urteil(args.eingabe, args.ausgabe, dokumenttyp=args.typ)
    print("Geschrieben: " + str(docx))
    if args.pdf:
        pdf = konvertiere_zu_pdf(docx)
        if pdf:
            print("PDF: " + str(pdf))
    return 0


if __name__ == "__main__":
    sys.exit(main())
