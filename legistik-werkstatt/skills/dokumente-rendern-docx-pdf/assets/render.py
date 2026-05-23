#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
render.py - Generiert deutsche Gesetzesentwurfsdokumente im offiziellen HdR-Layout.

Unterstuetzte Formate:
  - referentenentwurf      (Arial 11pt, ministerieller Hausstil, Bearbeitungsstand-Kopf)
  - bt-drucksache          (Times New Roman 11pt, Drucksachennummer + Wahlperiode in Kopfzeile)
  - formulierungshilfe     (Times New Roman 11pt, kurze Mantelseite)
  - synopse                (Dreispaltige Tabelle: geltend | Aenderung | Begruendung)
  - lesefassung            (konsolidierte Fassung, einspaltig)
  - kabinettsmappe         (Deckblatt fuer das Kabinett mit Tagesordnungspunkt)

Eingabe: Ordner mit metadaten.yaml, vorblatt.md, gesetzestext.md, begruendung-a.md,
begruendung-b.md, synopse.csv (optional anlagen/).

Ausgabe: DOCX im Output-Ordner. Optional PDF via soffice (LibreOffice).

Autor: Klotzkette - Plugin legistik-werkstatt
"""

import argparse
import csv
import os
import re
import shutil
import subprocess
import sys
from pathlib import Path

try:
    import yaml
except ImportError:
    yaml = None

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.section import WD_SECTION, WD_HEADER_FOOTER
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("FEHLER: python-docx fehlt. Bitte 'pip install python-docx pyyaml' ausfuehren.", file=sys.stderr)
    sys.exit(2)


# ---------- Schriftkonstanten nach Handbuch der Rechtsfoermlichkeit ----------

SCHRIFT_REFERENT = "Arial"
SCHRIFT_BTDRUCKSACHE = "Times New Roman"
SCHRIFTGROESSE_NORMAL = Pt(11)
SCHRIFTGROESSE_KLEIN = Pt(9)
SCHRIFTGROESSE_UEBERSCHRIFT = Pt(14)


def lade_metadaten(eingabe: Path) -> dict:
    meta_pfad = eingabe / "metadaten.yaml"
    if not meta_pfad.exists():
        return {
            "titel": "Entwurf eines Gesetzes",
            "kurztitel": "Beispielgesetz",
            "abkuerzung": "BeispG",
            "federfuehrung": "Bundesministerium der Justiz",
            "bearbeitungsstand": "TT.MM.JJJJ HH:MM",
            "drucksachennummer": "XX/YYYY",
            "wahlperiode": "21",
        }
    if yaml is None:
        # Minimaler Parser ohne pyyaml
        daten = {}
        for zeile in meta_pfad.read_text(encoding="utf-8").splitlines():
            if ":" in zeile and not zeile.strip().startswith("#"):
                k, v = zeile.split(":", 1)
                daten[k.strip()] = v.strip().strip('"').strip("'")
        return daten
    return yaml.safe_load(meta_pfad.read_text(encoding="utf-8"))


def setze_standardschrift(dokument: Document, schrift: str) -> None:
    style = dokument.styles["Normal"]
    style.font.name = schrift
    style.font.size = SCHRIFTGROESSE_NORMAL
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), schrift)
    rfonts.set(qn("w:hAnsi"), schrift)
    rfonts.set(qn("w:cs"), schrift)


def setze_seitenraender(dokument: Document) -> None:
    for section in dokument.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)


def fuege_sperrsatz_absatz(dokument: Document, text: str, *, fett: bool = True) -> None:
    """Fuegt eine Ueberschrift im Sperrsatz hinzu (Leerzeichen zwischen Buchstaben),
    wie es in BT-Drucksachen fuer Hauptueberschriften ueblich ist."""
    gesperrt = " ".join(list(text))
    p = dokument.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(gesperrt)
    r.bold = fett
    r.font.size = SCHRIFTGROESSE_UEBERSCHRIFT


def fuege_zentrierten_text(dokument: Document, text: str, *, fett: bool = False, groesse: Pt | None = None) -> None:
    p = dokument.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = fett
    if groesse:
        r.font.size = groesse


def setze_kopfzeile_referentenentwurf(dokument: Document, bearbeitungsstand: str) -> None:
    section = dokument.sections[0]
    header = section.header
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(f"Bearbeitungsstand: {bearbeitungsstand}")
    r.font.size = SCHRIFTGROESSE_KLEIN


def setze_kopfzeile_drucksache(dokument: Document, drucksachennummer: str, wahlperiode: str) -> None:
    section = dokument.sections[0]
    header = section.header
    # Tabelle mit 3 Spalten fuer links / mitte / rechts
    table = header.add_table(rows=1, cols=3, width=Cm(16.5))
    table.autofit = True
    tr = table.rows[0].cells
    p_l = tr[0].paragraphs[0]
    p_l.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_l.add_run(f"Drucksache {drucksachennummer}").font.size = SCHRIFTGROESSE_KLEIN
    p_m = tr[1].paragraphs[0]
    p_m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Seitennummerierung als Feld
    add_seitenzahl_feld(p_m)
    p_r = tr[2].paragraphs[0]
    p_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_r.add_run(f"Deutscher Bundestag - {wahlperiode}. Wahlperiode").font.size = SCHRIFTGROESSE_KLEIN


def add_seitenzahl_feld(paragraph) -> None:
    """Fuegt ein PAGE-Feld in das Absatz-Element ein, eingerahmt von Gedankenstrichen `- N -`."""
    paragraph.add_run("- ").font.size = SCHRIFTGROESSE_KLEIN
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run.font.size = SCHRIFTGROESSE_KLEIN
    paragraph.add_run(" -").font.size = SCHRIFTGROESSE_KLEIN


def lies_md_blocks(pfad: Path) -> list[tuple[str, str]]:
    """Liest Markdown und liefert Liste (Ebene-oder-'p', Inhalt).
    Ebenen: 'h1'..'h4' fuer Ueberschriften, 'p' fuer normale Absaetze, 'em' fuer kursiven Aenderungsbefehl."""
    bloecke = []
    if not pfad.exists():
        return bloecke
    for zeile in pfad.read_text(encoding="utf-8").splitlines():
        z = zeile.rstrip()
        if not z.strip():
            continue
        if z.startswith("#### "):
            bloecke.append(("h4", z[5:].strip()))
        elif z.startswith("### "):
            bloecke.append(("h3", z[4:].strip()))
        elif z.startswith("## "):
            bloecke.append(("h2", z[3:].strip()))
        elif z.startswith("# "):
            bloecke.append(("h1", z[2:].strip()))
        else:
            bloecke.append(("p", z.strip()))
    return bloecke


_INLINE_PATTERN = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")


def _schreibe_inline_text(absatz, text: str, *, default_italic: bool = False) -> None:
    """Zerlegt einen Markdown-Absatz in **fett**, *kursiv* und Normaltext-Runs."""
    pos = 0
    for m in _INLINE_PATTERN.finditer(text):
        if m.start() > pos:
            r = absatz.add_run(text[pos:m.start()])
            r.italic = default_italic
        token = m.group(0)
        if token.startswith("**"):
            r = absatz.add_run(token[2:-2])
            r.bold = True
            r.italic = default_italic
        else:
            r = absatz.add_run(token[1:-1])
            r.italic = True
        pos = m.end()
    if pos < len(text):
        r = absatz.add_run(text[pos:])
        r.italic = default_italic


def schreibe_bloecke(dokument: Document, bloecke: list[tuple[str, str]]) -> None:
    for ebene, inhalt in bloecke:
        if ebene == "h1":
            p = dokument.add_paragraph()
            r = p.add_run(inhalt)
            r.bold = True
            r.font.size = Pt(14)
        elif ebene == "h2":
            p = dokument.add_paragraph()
            r = p.add_run(inhalt)
            r.bold = True
            r.font.size = Pt(12)
        elif ebene == "h3":
            p = dokument.add_paragraph()
            r = p.add_run(inhalt)
            r.bold = True
        elif ebene == "h4":
            p = dokument.add_paragraph()
            r = p.add_run(inhalt)
            r.italic = True
        else:
            p = dokument.add_paragraph()
            # Heuristik: Wenn der ganze Absatz in Anfuehrungszeichen steht, ist es ein
            # Aenderungsbefehl - kursiv setzen.
            if inhalt.startswith('"') and inhalt.endswith('"'):
                _schreibe_inline_text(p, inhalt, default_italic=True)
            else:
                _schreibe_inline_text(p, inhalt, default_italic=False)


def baue_referentenentwurf(eingabe: Path, ausgabe: Path, meta: dict) -> Path:
    dokument = Document()
    setze_standardschrift(dokument, SCHRIFT_REFERENT)
    setze_seitenraender(dokument)
    setze_kopfzeile_referentenentwurf(dokument, meta.get("bearbeitungsstand", "TT.MM.JJJJ HH:MM"))

    # Titelseite
    fuege_zentrierten_text(dokument, "Referentenentwurf", fett=True, groesse=SCHRIFTGROESSE_UEBERSCHRIFT)
    federfuehrung = meta.get("federfuehrung", "Bundesministeriums der Justiz")
    # Genitiv-Korrektur: "Bundesministerium" -> "Bundesministeriums" wenn noetig
    genitiv = re.sub(r"^(Bundesministerium)(\s)", r"\1s\2", federfuehrung)
    fuege_zentrierten_text(dokument, f"des {genitiv}")
    dokument.add_paragraph()
    fuege_zentrierten_text(dokument, meta.get("titel", "Entwurf eines Gesetzes"), fett=True, groesse=SCHRIFTGROESSE_UEBERSCHRIFT)
    if meta.get("kurztitel"):
        kurz = meta.get("kurztitel")
        abk = meta.get("abkuerzung", "")
        if abk:
            fuege_zentrierten_text(dokument, f"({kurz} - {abk})", fett=True)
        else:
            fuege_zentrierten_text(dokument, f"({kurz})", fett=True)
    dokument.add_paragraph()
    fuege_zentrierten_text(dokument, "Vom ...")
    dokument.add_page_break()

    # Vorblatt
    dokument.add_heading("Vorblatt", level=1)
    schreibe_bloecke(dokument, lies_md_blocks(eingabe / "vorblatt.md"))
    dokument.add_page_break()

    # Gesetzestext
    dokument.add_heading("Gesetzestext", level=1)
    schreibe_bloecke(dokument, lies_md_blocks(eingabe / "gesetzestext.md"))
    dokument.add_page_break()

    # Begruendung
    dokument.add_heading("Begruendung", level=1)
    dokument.add_heading("A. Allgemeiner Teil", level=2)
    schreibe_bloecke(dokument, lies_md_blocks(eingabe / "begruendung-a.md"))
    dokument.add_heading("B. Besonderer Teil", level=2)
    schreibe_bloecke(dokument, lies_md_blocks(eingabe / "begruendung-b.md"))

    pfad = ausgabe / f"Referentenentwurf-{meta.get('abkuerzung', 'BeispG')}.docx"
    dokument.save(pfad)
    return pfad


def baue_bt_drucksache(eingabe: Path, ausgabe: Path, meta: dict) -> Path:
    dokument = Document()
    setze_standardschrift(dokument, SCHRIFT_BTDRUCKSACHE)
    setze_seitenraender(dokument)
    setze_kopfzeile_drucksache(dokument, meta.get("drucksachennummer", "XX/YYYY"), meta.get("wahlperiode", "21"))

    # Titelseite
    fuege_zentrierten_text(dokument, f"Deutscher Bundestag", fett=True, groesse=SCHRIFTGROESSE_UEBERSCHRIFT)
    fuege_zentrierten_text(dokument, f"Drucksache {meta.get('drucksachennummer', 'XX/YYYY')}", fett=True)
    fuege_zentrierten_text(dokument, f"{meta.get('wahlperiode', '21')}. Wahlperiode")
    dokument.add_paragraph()
    fuege_zentrierten_text(dokument, "Gesetzentwurf", fett=True, groesse=SCHRIFTGROESSE_UEBERSCHRIFT)
    fuege_zentrierten_text(dokument, "der Bundesregierung")
    dokument.add_paragraph()
    fuege_zentrierten_text(dokument, meta.get("titel", "Entwurf eines Gesetzes"), fett=True)
    if meta.get("kurztitel"):
        kurz = meta.get("kurztitel")
        abk = meta.get("abkuerzung", "")
        fuege_zentrierten_text(dokument, f"({kurz} - {abk})" if abk else f"({kurz})", fett=True)
    dokument.add_page_break()

    # Inhaltsuebersicht (Sperrsatz)
    fuege_sperrsatz_absatz(dokument, "Inhaltsuebersicht")
    dokument.add_paragraph()

    # Vorblatt
    schreibe_bloecke(dokument, lies_md_blocks(eingabe / "vorblatt.md"))
    dokument.add_page_break()

    # Anschreiben Bundeskanzler
    dokument.add_paragraph("BUNDESREPUBLIK DEUTSCHLAND").runs[0].bold = True
    dokument.add_paragraph("DER BUNDESKANZLER")
    dokument.add_paragraph()
    dokument.add_paragraph("An die Praesidentin des Deutschen Bundestages")
    dokument.add_paragraph()
    dokument.add_paragraph(
        f"Hiermit uebersende ich den von der Bundesregierung beschlossenen Entwurf eines "
        f"{meta.get('titel', 'Gesetzes')} mit Begruendung (Anlage 1) und Vorblatt. "
        f"Ich bitte, die Beschlussfassung des Deutschen Bundestages herbeizufuehren."
    )
    dokument.add_paragraph()
    dokument.add_paragraph("Federfuehrend ist das " + meta.get("federfuehrung", "Bundesministerium der Justiz") + ".")
    dokument.add_paragraph()
    dokument.add_paragraph("Mit freundlichen Gruessen")
    dokument.add_paragraph()
    dokument.add_paragraph(meta.get("kanzler", "Der Bundeskanzler"))
    dokument.add_page_break()

    # Gesetzestext
    fuege_sperrsatz_absatz(dokument, "Gesetzestext")
    schreibe_bloecke(dokument, lies_md_blocks(eingabe / "gesetzestext.md"))
    dokument.add_page_break()

    # Begruendung
    fuege_sperrsatz_absatz(dokument, "Begruendung")
    dokument.add_heading("A. Allgemeiner Teil", level=2)
    schreibe_bloecke(dokument, lies_md_blocks(eingabe / "begruendung-a.md"))
    dokument.add_heading("B. Besonderer Teil", level=2)
    schreibe_bloecke(dokument, lies_md_blocks(eingabe / "begruendung-b.md"))

    pfad = ausgabe / f"BT-Drucksache-{meta.get('drucksachennummer', 'XX-YYYY').replace('/', '-')}-{meta.get('abkuerzung', 'BeispG')}.docx"
    dokument.save(pfad)
    return pfad


def baue_formulierungshilfe(eingabe: Path, ausgabe: Path, meta: dict) -> Path:
    dokument = Document()
    setze_standardschrift(dokument, SCHRIFT_BTDRUCKSACHE)
    setze_seitenraender(dokument)
    setze_kopfzeile_referentenentwurf(dokument, meta.get("bearbeitungsstand", "TT.MM.JJJJ HH:MM"))

    fuege_zentrierten_text(dokument, "Formulierungshilfe der Bundesregierung", fett=True, groesse=SCHRIFTGROESSE_UEBERSCHRIFT)
    fuege_zentrierten_text(dokument, f"fuer die Fraktionen des Deutschen Bundestages")
    dokument.add_paragraph()
    fuege_zentrierten_text(dokument, meta.get("titel", "Entwurf eines Gesetzes"), fett=True)
    dokument.add_paragraph()
    fuege_zentrierten_text(dokument, "(Aenderungsantrag zum Gesetzentwurf der Fraktionen ...)")
    dokument.add_page_break()

    dokument.add_heading("I. Zielsetzung", level=2)
    schreibe_bloecke(dokument, lies_md_blocks(eingabe / "vorblatt.md"))
    dokument.add_heading("II. Aenderungstext", level=2)
    schreibe_bloecke(dokument, lies_md_blocks(eingabe / "gesetzestext.md"))
    dokument.add_heading("III. Begruendung", level=2)
    schreibe_bloecke(dokument, lies_md_blocks(eingabe / "begruendung-b.md"))

    pfad = ausgabe / f"Formulierungshilfe-{meta.get('abkuerzung', 'BeispG')}.docx"
    dokument.save(pfad)
    return pfad


def baue_synopse(eingabe: Path, ausgabe: Path, meta: dict) -> Path:
    dokument = Document()
    setze_standardschrift(dokument, SCHRIFT_REFERENT)
    setze_seitenraender(dokument)

    fuege_zentrierten_text(dokument, "Synopse", fett=True, groesse=SCHRIFTGROESSE_UEBERSCHRIFT)
    fuege_zentrierten_text(dokument, meta.get("titel", "Entwurf eines Gesetzes"))
    dokument.add_paragraph()

    csv_pfad = eingabe / "synopse.csv"
    if not csv_pfad.exists():
        dokument.add_paragraph("Keine synopse.csv im Eingabeordner gefunden.")
    else:
        with csv_pfad.open(encoding="utf-8", newline="") as fh:
            reader = csv.reader(fh, delimiter=";")
            rows = list(reader)
        if rows:
            tabelle = dokument.add_table(rows=1, cols=3)
            tabelle.style = "Light Grid Accent 1"
            hdr = tabelle.rows[0].cells
            hdr[0].text = rows[0][0] if len(rows[0]) > 0 else "Geltendes Recht"
            hdr[1].text = rows[0][1] if len(rows[0]) > 1 else "Aenderung"
            hdr[2].text = rows[0][2] if len(rows[0]) > 2 else "Begruendung"
            for r in rows[1:]:
                cells = tabelle.add_row().cells
                cells[0].text = r[0] if len(r) > 0 else ""
                cells[1].text = r[1] if len(r) > 1 else ""
                cells[2].text = r[2] if len(r) > 2 else ""

    pfad = ausgabe / f"Synopse-{meta.get('abkuerzung', 'BeispG')}.docx"
    dokument.save(pfad)
    return pfad


def baue_lesefassung(eingabe: Path, ausgabe: Path, meta: dict) -> Path:
    dokument = Document()
    setze_standardschrift(dokument, SCHRIFT_REFERENT)
    setze_seitenraender(dokument)

    fuege_zentrierten_text(dokument, "Lesefassung (konsolidiert)", fett=True, groesse=SCHRIFTGROESSE_UEBERSCHRIFT)
    fuege_zentrierten_text(dokument, meta.get("titel", "Entwurf eines Gesetzes"))
    fuege_zentrierten_text(dokument, "nach Inkrafttreten der vorgeschlagenen Aenderungen")
    dokument.add_paragraph()
    schreibe_bloecke(dokument, lies_md_blocks(eingabe / "lesefassung.md"))

    pfad = ausgabe / f"Lesefassung-{meta.get('abkuerzung', 'BeispG')}.docx"
    dokument.save(pfad)
    return pfad


def baue_kabinettsmappe(eingabe: Path, ausgabe: Path, meta: dict) -> Path:
    dokument = Document()
    setze_standardschrift(dokument, SCHRIFT_REFERENT)
    setze_seitenraender(dokument)

    fuege_zentrierten_text(dokument, "Kabinettsmappe", fett=True, groesse=SCHRIFTGROESSE_UEBERSCHRIFT)
    dokument.add_paragraph()
    dokument.add_paragraph(f"Tagesordnungspunkt: {meta.get('top', 'TBA')}")
    dokument.add_paragraph(f"Federfuehrung: {meta.get('federfuehrung', 'Bundesministerium der Justiz')}")
    dokument.add_paragraph(f"Bearbeitungsstand: {meta.get('bearbeitungsstand', 'TT.MM.JJJJ')}")
    dokument.add_paragraph()
    dokument.add_paragraph("Anlagen:")
    dokument.add_paragraph("Anlage 1 - Referentenentwurf")
    dokument.add_paragraph("Anlage 2 - Vorblatt mit Erfuellungsaufwand")
    dokument.add_paragraph("Anlage 3 - Stellungnahme Normenkontrollrat")
    dokument.add_paragraph("Anlage 4 - Sprechzettel des Ressorts")

    pfad = ausgabe / f"Kabinettsmappe-{meta.get('abkuerzung', 'BeispG')}.docx"
    dokument.save(pfad)
    return pfad


def konvertiere_zu_pdf(docx_pfad: Path) -> Path | None:
    """Versucht, das DOCX via soffice headless zu PDF zu konvertieren."""
    if shutil.which("soffice") is None and shutil.which("libreoffice") is None:
        return None
    binary = shutil.which("soffice") or shutil.which("libreoffice")
    try:
        subprocess.run(
            [binary, "--headless", "--convert-to", "pdf", "--outdir", str(docx_pfad.parent), str(docx_pfad)],
            check=True, capture_output=True, timeout=120,
        )
    except (subprocess.CalledProcessError, subprocess.TimeoutExpired):
        return None
    pdf_pfad = docx_pfad.with_suffix(".pdf")
    return pdf_pfad if pdf_pfad.exists() else None


RENDERER = {
    "referentenentwurf": baue_referentenentwurf,
    "bt-drucksache": baue_bt_drucksache,
    "formulierungshilfe": baue_formulierungshilfe,
    "synopse": baue_synopse,
    "lesefassung": baue_lesefassung,
    "kabinettsmappe": baue_kabinettsmappe,
}


def main() -> int:
    parser = argparse.ArgumentParser(description="HdR-konformer Renderer fuer Legistik-Dokumente.")
    parser.add_argument("--format", required=True, choices=list(RENDERER.keys()))
    parser.add_argument("--eingabe", required=True, help="Eingabeordner mit metadaten.yaml und md-Bausteinen.")
    parser.add_argument("--ausgabe", required=True, help="Zielordner fuer DOCX (und ggf. PDF).")
    parser.add_argument("--pdf", action="store_true", help="Zusaetzlich PDF erzeugen, falls soffice verfuegbar.")
    args = parser.parse_args()

    eingabe = Path(args.eingabe)
    ausgabe = Path(args.ausgabe)
    ausgabe.mkdir(parents=True, exist_ok=True)
    if not eingabe.exists():
        print(f"FEHLER: Eingabeordner {eingabe} existiert nicht.", file=sys.stderr)
        return 1

    meta = lade_metadaten(eingabe)
    pfad = RENDERER[args.format](eingabe, ausgabe, meta)
    print(f"Geschrieben: {pfad}")

    if args.pdf:
        pdf = konvertiere_zu_pdf(pfad)
        if pdf:
            print(f"PDF geschrieben: {pdf}")
        else:
            print("Hinweis: PDF-Konvertierung uebersprungen (kein soffice/libreoffice gefunden).", file=sys.stderr)

    return 0


if __name__ == "__main__":
    sys.exit(main())
