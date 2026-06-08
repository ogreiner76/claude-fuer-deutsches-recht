#!/usr/bin/env python3
"""Baut für jede Testakte ein 'gesamt-pdf/<name>_gesamt.pdf', das die
exportfaehigen Aktenstücke (MD/TXT/EML/CSV/XLSX/DOCX/Bilder/PDF) in ein
einziges, sauber gerendertes Dokument mit Dateigrenzen und Seitenzahlen
zusammenfasst.

Aufruf:
  python3 scripts/build-testakte-gesamt-pdf.py                 # alle Testakten
  python3 scripts/build-testakte-gesamt-pdf.py <name1> <name2>  # gezielt
"""

from __future__ import annotations

import io
import re
import sys
import csv
from email import policy
from email.parser import BytesParser
from pathlib import Path

# Drittabhaengigkeiten
from openpyxl import load_workbook
from pypdf import PdfReader, PdfWriter
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib.colors import HexColor, black
from reportlab.platypus import (
    SimpleDocTemplate,
    Image as RLImage,
    Paragraph,
    Spacer,
    PageBreak,
    Table,
    TableStyle,
)
from reportlab.lib.utils import ImageReader
from reportlab.lib.enums import TA_LEFT
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from testakte_file_filter import include_in_working_dump

# DOCX
try:
    from docx import Document
except ImportError:
    Document = None  # type: ignore

REPO_ROOT = Path(__file__).resolve().parent.parent
TESTAKTEN = REPO_ROOT / "testakten"

# Design
TEAL = HexColor("#01696F")
MUTED = HexColor("#7A7974")
BORDER = HexColor("#D4D1CA")
SURFACE = HexColor("#F7F6F2")

# Font: System-Helvetica als Fallback. Inter waere schoener, aber wir verzichten
# auf Netzwerk-Downloads, damit das Skript offline laeuft.
FONT_REG = "Helvetica"
FONT_BOLD = "Helvetica-Bold"

styles = getSampleStyleSheet()
s_cover_label = ParagraphStyle(
    "CoverLabel",
    fontName=FONT_REG, fontSize=14, leading=18,
    textColor=MUTED, spaceAfter=6,
)
s_cover_title = ParagraphStyle(
    "CoverTitle",
    fontName=FONT_BOLD, fontSize=28, leading=34,
    textColor=TEAL, alignment=TA_LEFT, spaceAfter=14,
)
s_cover_sub = ParagraphStyle(
    "CoverSub",
    fontName=FONT_REG, fontSize=12, leading=16,
    textColor=black, spaceAfter=4,
)
s_cover_meta = ParagraphStyle(
    "CoverMeta",
    fontName=FONT_REG, fontSize=9, leading=12,
    textColor=MUTED, spaceAfter=3,
)
s_h1 = ParagraphStyle(
    "H1", parent=styles["Heading1"],
    fontName=FONT_BOLD, fontSize=18, leading=22, textColor=TEAL,
    spaceBefore=18, spaceAfter=8,
)
s_h2 = ParagraphStyle(
    "H2", parent=styles["Heading2"],
    fontName=FONT_BOLD, fontSize=14, leading=18, textColor=black,
    spaceBefore=12, spaceAfter=6,
)
s_h3 = ParagraphStyle(
    "H3", parent=styles["Heading3"],
    fontName=FONT_BOLD, fontSize=11, leading=14, textColor=black,
    spaceBefore=8, spaceAfter=4,
)
s_body = ParagraphStyle(
    "Body", parent=styles["BodyText"],
    fontName=FONT_REG, fontSize=10, leading=14, textColor=black,
    spaceAfter=6,
)
s_meta = ParagraphStyle(
    "Meta", parent=styles["BodyText"],
    fontName=FONT_REG, fontSize=9, leading=12, textColor=MUTED,
    spaceAfter=4,
)
s_partlabel = ParagraphStyle(
    "PartLabel", parent=styles["BodyText"],
    fontName=FONT_BOLD, fontSize=11, leading=14, textColor=MUTED,
    spaceAfter=2,
)

# Reihenfolge der Datei-Typen im Gesamt-PDF
TYPE_ORDER = ["md", "txt", "eml", "csv", "xlsx", "docx", "image", "pdf"]
IMAGE_EXTS = {"jpg", "jpeg", "png"}
TYPE_LABEL = {
    "md": "Aktenstücke",
    "txt": "Notizen und Textdateien",
    "eml": "E-Mails",
    "csv": "CSV-Tabellen",
    "xlsx": "Excel-Tabellen",
    "docx": "Word-Dokumente",
    "image": "Bildanlagen und Screenshots",
    "pdf": "PDF-Anhänge (Originaldokumente)",
}


def escape(s: str) -> str:
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def md_to_flowables(md_text: str) -> list:
    out: list = []
    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue
        if line.startswith("# "):
            out.append(Paragraph(escape(line[2:].strip()), s_h1))
            i += 1
            continue
        if line.startswith("## "):
            out.append(Paragraph(escape(line[3:].strip()), s_h2))
            i += 1
            continue
        if line.startswith("### "):
            out.append(Paragraph(escape(line[4:].strip()), s_h3))
            i += 1
            continue
        if line.startswith("---"):
            out.append(Spacer(1, 6))
            i += 1
            continue
        # Tabelle?
        if (
            line.startswith("|")
            and i + 1 < len(lines)
            and re.match(r"^\|[\s\-:|]+\|$", lines[i + 1])
        ):
            header = [c.strip() for c in line.strip("|").split("|")]
            i += 2
            rows = [header]
            while i < len(lines) and lines[i].startswith("|"):
                cells = [c.strip() for c in lines[i].strip("|").split("|")]
                rows.append(cells)
                i += 1
            col_count = max(1, len(header))
            avail_width = 16 * cm
            col_widths = [avail_width / col_count] * col_count
            if col_count >= 2:
                col_widths[0] = min(4 * cm, avail_width / col_count * 1.5)
                rest = (avail_width - col_widths[0]) / (col_count - 1)
                for k in range(1, col_count):
                    col_widths[k] = rest
            tbl_data = []
            for r in rows:
                tbl_data.append([Paragraph(escape(c), s_body) for c in r])
            tbl = Table(tbl_data, colWidths=col_widths, repeatRows=1)
            tbl.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), SURFACE),
                        ("FONTNAME", (0, 0), (-1, 0), FONT_BOLD),
                        ("BOX", (0, 0), (-1, -1), 0.4, BORDER),
                        ("INNERGRID", (0, 0), (-1, -1), 0.3, BORDER),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 4),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                        ("TOPPADDING", (0, 0), (-1, -1), 3),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                    ]
                )
            )
            out.append(tbl)
            out.append(Spacer(1, 6))
            continue
        if line.startswith("- ") or line.startswith("* "):
            text = line[2:].strip()
            text = re.sub(r"\*\*([^*]+)\*\*", r"<b>\1</b>", text)
            out.append(Paragraph("• " + _inline_markup(text), s_body))
            i += 1
            continue
        if re.match(r"^\d+\.\s", line):
            text = re.sub(r"\*\*([^*]+)\*\*", r"<b>\1</b>", line)
            out.append(Paragraph(_inline_markup(text), s_body))
            i += 1
            continue
        # Sammle normalen Absatz bis zur naechsten Leerzeile/Sondersyntax
        block = [line]
        j = i + 1
        while (
            j < len(lines)
            and lines[j].strip()
            and not lines[j].startswith(("#", "-", "*", "|", "---"))
            and not re.match(r"^\d+\.\s", lines[j])
        ):
            block.append(lines[j].rstrip())
            j += 1
        text = " ".join(block).strip()
        text = re.sub(r"\*\*([^*]+)\*\*", r"<b>\1</b>", text)
        text = re.sub(r"`([^`]+)`", r"<font face='Courier'>\1</font>", text)
        out.append(Paragraph(_inline_markup(text), s_body))
        i = j
    return out


def _inline_markup(s: str) -> str:
    """Escape minimal, lasse erlaubte Inline-Tags."""
    s = s.replace("&", "&amp;")
    # Erlaubte Tags wieder herstellen
    s = re.sub(r"&lt;(/?(?:b|i|sub|sup))&gt;", r"<\1>", s)
    s = s.replace("&lt;font face='Courier'&gt;", "<font face='Courier'>")
    s = s.replace("&lt;/font&gt;", "</font>")
    # Eckige Klammern in normalem Text: behalten, aber nicht als Tag
    return s


def txt_to_flowables(text: str) -> list:
    out = []
    for para in text.split("\n\n"):
        para = para.strip()
        if not para:
            continue
        # Zeilenumbrueche im Absatz erhalten
        out.append(Paragraph(escape(para).replace("\n", "<br/>"), s_body))
    return out


def eml_to_flowables(path: Path) -> list:
    out = []
    try:
        with open(path, "rb") as f:
            msg = BytesParser(policy=policy.default).parse(f)
        headers = [
            ("Von", msg.get("From", "")),
            ("An", msg.get("To", "")),
            ("Datum", msg.get("Date", "")),
            ("Betreff", msg.get("Subject", "")),
        ]
        body_part = msg.get_body(preferencelist=("plain", "html"))
        body = body_part.get_content() if body_part else ""
    except Exception as e:
        out.append(Paragraph(f"<i>E-Mail konnte nicht gelesen werden: {escape(str(e))}</i>", s_meta))
        return out

    rows = [
        [Paragraph(label, s_meta), Paragraph(escape(value), s_meta)]
        for label, value in headers
    ]
    tbl = Table(rows, colWidths=[2.5 * cm, 13.5 * cm])
    tbl.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, -1), SURFACE),
                ("BOX", (0, 0), (-1, -1), 0.3, BORDER),
                ("INNERGRID", (0, 0), (-1, -1), 0.2, BORDER),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]
        )
    )
    out.append(tbl)
    out.append(Spacer(1, 6))
    out.extend(txt_to_flowables(body))
    return out


def csv_to_flowables(path: Path) -> list:
    out = []
    try:
        with open(path, encoding="utf-8") as f:
            reader = csv.reader(f)
            rows = list(reader)
    except UnicodeDecodeError:
        with open(path, encoding="latin-1") as f:
            reader = csv.reader(f)
            rows = list(reader)
    except Exception as e:
        out.append(Paragraph(f"<i>CSV konnte nicht gelesen werden: {escape(str(e))}</i>", s_meta))
        return out

    if not rows:
        return out
    max_cols = max(len(r) for r in rows)
    rows = [r + [""] * (max_cols - len(r)) for r in rows]
    out.extend(_render_table(rows, header=True))
    return out


def xlsx_to_flowables(path: Path) -> list:
    out = []
    try:
        wb = load_workbook(path, data_only=True)
    except Exception as e:
        out.append(Paragraph(f"<i>XLSX konnte nicht gelesen werden: {escape(str(e))}</i>", s_meta))
        return out
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        out.append(Paragraph(f"Tabellenblatt: {escape(sheet_name)}", s_h3))
        rows = []
        for row in ws.iter_rows(values_only=True):
            rows.append(
                [_format_cell(c) for c in row]
            )
        if not rows:
            continue
        # Leere Zeilen/Spalten am Ende abschneiden
        while rows and not any(c.strip() for c in rows[-1]):
            rows.pop()
        if not rows:
            continue
        max_cols = max(len(r) for r in rows)
        if max_cols == 0:
            continue
        # Hinten leere Spalten abschneiden
        while max_cols > 0 and all(
            (len(r) <= max_cols - 1) or (not r[max_cols - 1].strip()) for r in rows
        ):
            max_cols -= 1
        if max_cols == 0:
            continue
        rows = [r[:max_cols] + [""] * (max_cols - len(r[:max_cols])) for r in rows]
        out.extend(_render_table(rows, header=True))
        out.append(Spacer(1, 6))
    return out


def _format_cell(c) -> str:
    if c is None:
        return ""
    if isinstance(c, float):
        if c == int(c):
            return str(int(c))
        return f"{c:.4f}".rstrip("0").rstrip(".")
    return str(c)


# Maximalzeichen pro Zelle, ab denen die Tabelle nicht mehr als Table gerendert wird,
# sondern als sequentielle Absatzfolge (verhindert ReportLab-Overflow).
_MAX_CELL_CHARS = 1200


def _split_long_text(text: str, chunk: int = 800) -> list:
    """Schneidet sehr langen Text an Absatz- oder Satzgrenzen in Stuecke."""
    text = text.replace("\r", "")
    if len(text) <= chunk:
        return [text]
    # Erst Absaetze probieren
    paras = [p for p in text.split("\n") if p.strip()]
    if any(len(p) > chunk for p in paras):
        # Weiter an Saetzen schneiden
        out = []
        for p in paras:
            if len(p) <= chunk:
                out.append(p)
                continue
            buf = ""
            for sent in p.replace("; ", "; |").replace(". ", ". |").split("|"):
                if len(buf) + len(sent) > chunk and buf:
                    out.append(buf)
                    buf = sent
                else:
                    buf += sent
            if buf:
                out.append(buf)
        return out
    return paras


def _render_table(rows: list, header: bool = False) -> list:
    """Rendert eine Tabelle. Falls Zellen zu lang werden, faellt es auf eine
    sequentielle Absatzdarstellung zurueck (Reihe fuer Reihe), damit ReportLab
    keine Overflow-Fehler wirft."""
    max_cell_len = max((len(c) for r in rows for c in r), default=0)
    max_cols_in_table = max((len(r) for r in rows), default=0)
    # Bei sehr breiten Tabellen (>12 Spalten) faellt es sequentiell zurueck,
    # weil die Spaltenbreite sonst kleiner ist als die kleinste Wortbreite
    # und ReportLab Cell-Overflow-Fehler wirft.
    if max_cell_len > _MAX_CELL_CHARS or max_cols_in_table > 12:
        out = []
        header_row = rows[0] if header else None
        body_rows = rows[1:] if header else rows
        for ri, r in enumerate(body_rows):
            if header_row:
                # Reihen-Trennlinie + Spaltenkopf pro Zelle
                for ci, cell in enumerate(r):
                    label = header_row[ci] if ci < len(header_row) else f"Spalte {ci+1}"
                    if label.strip():
                        out.append(Paragraph(f"<b>{escape(label)}</b>", s_meta))
                    for chunk in _split_long_text(cell):
                        out.append(Paragraph(escape(chunk), s_body))
            else:
                for ci, cell in enumerate(r):
                    for chunk in _split_long_text(cell):
                        out.append(Paragraph(escape(chunk), s_body))
            out.append(Spacer(1, 4))
            out.append(HRFlowable(width="100%", thickness=0.2, color=BORDER))
            out.append(Spacer(1, 4))
        return out

    max_cols = max(len(r) for r in rows)
    avail_width = 16 * cm
    col_widths = [avail_width / max_cols] * max_cols
    data = [
        [Paragraph(escape(c), s_meta) for c in r]
        for r in rows
    ]
    tbl = Table(data, colWidths=col_widths, repeatRows=1 if header else 0, splitByRow=1)
    cmds = [
        ("BOX", (0, 0), (-1, -1), 0.3, BORDER),
        ("INNERGRID", (0, 0), (-1, -1), 0.2, BORDER),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 3),
        ("RIGHTPADDING", (0, 0), (-1, -1), 3),
        ("TOPPADDING", (0, 0), (-1, -1), 2),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
    ]
    if header:
        cmds.insert(0, ("BACKGROUND", (0, 0), (-1, 0), SURFACE))
        cmds.insert(1, ("FONTNAME", (0, 0), (-1, 0), FONT_BOLD))
    tbl.setStyle(TableStyle(cmds))
    return [tbl]


def docx_to_flowables(path: Path) -> list:
    out = []
    if Document is None:
        out.append(Paragraph("<i>python-docx nicht installiert, Inhalt wird uebersprungen.</i>", s_meta))
        return out
    try:
        doc = Document(str(path))
    except Exception as e:
        out.append(Paragraph(f"<i>DOCX konnte nicht gelesen werden: {escape(str(e))}</i>", s_meta))
        return out
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style else ""
        if style.startswith("Heading 1"):
            out.append(Paragraph(escape(text), s_h2))
        elif style.startswith("Heading 2"):
            out.append(Paragraph(escape(text), s_h3))
        elif style.startswith("Heading"):
            out.append(Paragraph(escape(text), s_h3))
        else:
            out.append(Paragraph(escape(text), s_body))
    for table in doc.tables:
        rows = []
        for r in table.rows:
            rows.append([c.text.strip() for c in r.cells])
        if rows:
            out.extend(_render_table(rows, header=True))
            out.append(Spacer(1, 6))
    return out


def image_to_flowables(path: Path) -> list:
    out = []
    try:
        width, height = ImageReader(str(path)).getSize()
        max_width = 16 * cm
        max_height = 22 * cm
        scale = min(max_width / width, max_height / height, 1)
        img = RLImage(str(path), width=width * scale, height=height * scale)
        out.append(img)
        out.append(Spacer(1, 4))
        out.append(Paragraph(f"Bilddatei: {escape(path.name)}", s_meta))
    except Exception as e:
        out.append(Paragraph(f"<i>Bild konnte nicht gerendert werden: {escape(str(e))}</i>", s_meta))
    return out


def header_footer_factory(testakte_name: str):
    def hf(canv: canvas.Canvas, doc) -> None:
        canv.saveState()
        canv.setFont(FONT_REG, 8)
        canv.setFillColor(MUTED)
        canv.drawString(2 * cm, 1.2 * cm, f"Arbeitsakte: {testakte_name}")
        canv.drawRightString(19 * cm, 1.2 * cm, f"Seite {doc.page}")
        canv.setStrokeColor(BORDER)
        canv.setLineWidth(0.3)
        canv.line(2 * cm, 1.6 * cm, 19 * cm, 1.6 * cm)
        canv.restoreState()

    return hf


def no_header_footer(canv: canvas.Canvas, doc) -> None:
    return None


def build_cover(_name: str, _readme_summary: str | None, h1: str | None = None) -> list:
    # Historisch gab es hier ein Titelblatt. Das Gesamt-PDF soll jetzt direkt
    # mit dem ersten Aktenstück beginnen; die Funktion bleibt als Kompatibilitaet
    # fuer aeltere Aufrufe erhalten.
    return []


def extract_readme_summary(readme_path: Path) -> tuple[str | None, str | None]:
    """Liest aus der README den H1-Titel und einen kurzen beschreibenden Absatz.

    Der beschreibende Absatz wird absichtlich erst gesucht, NACHDEM Download-Bloecke
    und Aktenstruktur-Blocks uebersprungen wurden, damit nicht der ZIP-Hinweis
    auf dem Cover landet.
    """
    if not readme_path.is_file():
        return None, None
    text = readme_path.read_text(encoding="utf-8")
    # H1
    h1_match = re.search(r"^#\s+(.+?)\s*$", text, flags=re.MULTILINE)
    h1 = h1_match.group(1).strip() if h1_match else None

    # Suche eine Sektion mit beschreibendem Inhalt: 'Kurzbild', 'Worum',
    # 'Sachverhalt', 'Ueberblick', 'Mandat', 'Fall' o.ae.
    section_pattern = re.compile(
        r"^##[^\n]*?(?:kurzbild|worum geht|sachverhalt|ueberblick|\u00fcberblick|mandat|fall|der fall|akte|kontext|ausgangslage|ausgangs|zweck|szenario|idee|einsatz|\u00fcbersicht|uebersicht|verfahrenseckdaten|aktenkern|aktenbestand|mandantenkonstellation|politische vorgabe|enthaltene arbeitsdateien|dateien)[^\n]*\n([\s\S]*?)(?=^## |\Z)",
        re.IGNORECASE | re.MULTILINE,
    )
    m = section_pattern.search(text)
    candidate_text = m.group(1) if m else text
    for para in candidate_text.split("\n\n"):
        para = para.strip()
        if not para:
            continue
        if para.startswith(("#", "-", "*", "|", "<!--", "```")):
            continue
        # Download-/ZIP-Hinweise ueberspringen
        lower = para.lower()
        if any(
            kw in lower
            for kw in (
                "zip-datei",
                "zip datei",
                "direkt-download",
                "als zip",
                "github-release",
                "github release",
                "download",
            )
        ):
            continue
        para = re.sub(r"\*\*([^*]+)\*\*", r"\1", para)
        para = re.sub(r"\s+", " ", para)
        return h1, para[:400]
    return h1, None


def collect_files(testakte_dir: Path) -> dict[str, list[Path]]:
    files_by_type: dict[str, list[Path]] = {t: [] for t in TYPE_ORDER}
    for f in testakte_dir.rglob("*"):
        if not include_in_working_dump(f, testakte_dir):
            continue
        ext = f.suffix.lower().lstrip(".")
        if ext in IMAGE_EXTS:
            files_by_type["image"].append(f)
            continue
        if ext not in TYPE_ORDER:
            continue
        files_by_type[ext].append(f)
    for t in files_by_type:
        files_by_type[t].sort(key=lambda p: str(p.relative_to(testakte_dir)).lower())
    return files_by_type


def build_text_pdf(testakte_dir: Path, files: dict[str, list[Path]], cover: list, tmp_path: Path) -> tuple[bool, list[Path]]:
    """Baut den Text-Teil als PDF, sammelt PDF-Anhaenge separat."""
    doc = SimpleDocTemplate(
        str(tmp_path),
        pagesize=A4,
        leftMargin=2 * cm, rightMargin=2 * cm,
        topMargin=2 * cm, bottomMargin=2 * cm,
        title=f"Arbeitsakte {testakte_dir.name}",
        author="Kanzleiakte",
    )
    flow = list(cover)

    pdf_attachments: list[Path] = []
    for t in TYPE_ORDER:
        if not files[t]:
            continue
        if t == "pdf":
            # PDFs werden separat angehaengt (Original-Layout bewahren)
            pdf_attachments = files[t]
            continue
        for f in files[t]:
            rel = f.relative_to(testakte_dir)
            flow.append(Paragraph(f"<b>Datei:</b> {escape(str(rel))}", s_meta))
            flow.append(Spacer(1, 4))
            try:
                if t == "md":
                    flow.extend(md_to_flowables(f.read_text(encoding="utf-8", errors="replace")))
                elif t == "txt":
                    flow.extend(txt_to_flowables(f.read_text(encoding="utf-8", errors="replace")))
                elif t == "eml":
                    flow.extend(eml_to_flowables(f))
                elif t == "csv":
                    flow.extend(csv_to_flowables(f))
                elif t == "xlsx":
                    flow.extend(xlsx_to_flowables(f))
                elif t == "docx":
                    flow.extend(docx_to_flowables(f))
                elif t == "image":
                    flow.extend(image_to_flowables(f))
            except Exception as e:
                flow.append(Paragraph(f"<i>Inhalt konnte nicht gerendert werden: {escape(str(e))}</i>", s_meta))
            flow.append(Spacer(1, 14))
        if flow:
            flow.append(PageBreak())

    if not flow:
        flow.append(Paragraph("Dateiablage: Original-PDFs folgen.", s_meta))

    hf = header_footer_factory(testakte_dir.name)
    try:
        doc.build(flow, onFirstPage=no_header_footer, onLaterPages=hf)
    except Exception as e:
        print(f"  FEHLER beim Bauen: {e}")
        return False, pdf_attachments
    return True, pdf_attachments


def append_pdf_with_separator(writer: PdfWriter, label: str, pdf_path: Path, testakte_name: str) -> None:
    sep = io.BytesIO()
    c = canvas.Canvas(sep, pagesize=A4)
    c.setTitle(label)
    c.setAuthor("Kanzleiakte")
    c.setFont(FONT_BOLD, 14)
    c.setFillColor(TEAL)
    c.drawString(2 * cm, 25 * cm, "Datei")
    c.setFont(FONT_REG, 9)
    c.setFillColor(MUTED)
    c.drawString(2 * cm, 24.2 * cm, label)
    c.setStrokeColor(BORDER)
    c.setLineWidth(0.3)
    c.line(2 * cm, 1.6 * cm, 19 * cm, 1.6 * cm)
    c.setFont(FONT_REG, 8)
    c.drawString(2 * cm, 1.2 * cm, testakte_name)
    c.showPage()
    c.save()
    sep.seek(0)
    for p in PdfReader(sep).pages:
        writer.add_page(p)
    try:
        for p in PdfReader(str(pdf_path)).pages:
            writer.add_page(p)
    except Exception as e:
        # PDF defekt oder verschluesselt -> Hinweisseite einfuegen
        sep2 = io.BytesIO()
        c2 = canvas.Canvas(sep2, pagesize=A4)
        c2.setFont(FONT_REG, 10)
        c2.drawString(2 * cm, 25 * cm, f"PDF konnte nicht eingebunden werden: {e}")
        c2.showPage()
        c2.save()
        sep2.seek(0)
        for p in PdfReader(sep2).pages:
            writer.add_page(p)


def build_gesamt_pdf(testakte_dir: Path) -> tuple[str, str]:
    """Gibt (status, info) zurueck. status in {ok, skip, error}."""
    name = testakte_dir.name
    out_dir = testakte_dir / "gesamt-pdf"
    out_dir.mkdir(exist_ok=True)
    out_path = out_dir / f"{name}_gesamt.pdf"

    files = collect_files(testakte_dir)
    total_files = sum(len(v) for v in files.values())
    if total_files == 0:
        return "skip", "keine Quelldateien"

    cover: list = []

    tmp_text = Path(f"/tmp/_gesamt_text_{name}.pdf")
    ok, pdf_attachments = build_text_pdf(testakte_dir, files, cover, tmp_text)
    if not ok:
        return "error", "Text-PDF konnte nicht erzeugt werden"

    writer = PdfWriter()
    try:
        for page in PdfReader(str(tmp_text)).pages:
            writer.add_page(page)
    except Exception as e:
        return "error", f"Text-PDF nicht lesbar: {e}"

    for pdf in pdf_attachments:
        rel = pdf.relative_to(testakte_dir)
        label = f"PDF-Anhang: {rel}"
        append_pdf_with_separator(writer, label, pdf, name)

    writer.add_metadata(
        {
            "/Title": f"Arbeitsakte {name}",
            "/Author": "Kanzleiakte",
            "/Subject": "Gesamtakte",
        }
    )
    with open(out_path, "wb") as f:
        writer.write(f)
    size_kb = out_path.stat().st_size / 1024
    try:
        tmp_text.unlink()
    except Exception:
        pass
    return "ok", f"{out_path.relative_to(REPO_ROOT)} ({size_kb:.0f} KB, {total_files} Quelldateien)"


def main() -> None:
    targets = sys.argv[1:]
    all_dirs = sorted([d for d in TESTAKTEN.iterdir() if d.is_dir()])
    if targets:
        all_dirs = [d for d in all_dirs if d.name in targets]
    print(f"Verarbeite {len(all_dirs)} Testakten")
    print()
    counts = {"ok": 0, "skip": 0, "error": 0}
    for d in all_dirs:
        status, info = build_gesamt_pdf(d)
        counts[status] += 1
        sigil = {"ok": "OK ", "skip": "SK ", "error": "ERR"}[status]
        print(f"  {sigil} {d.name}: {info}")
    print()
    print(f"Fertig: {counts['ok']} OK, {counts['skip']} skip, {counts['error']} Fehler")


if __name__ == "__main__":
    main()
